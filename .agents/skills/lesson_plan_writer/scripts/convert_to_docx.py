import sys
import os
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def create_element(name):
    return OxmlElement(name)

def set_cell_background(cell, color_hex):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def add_callout_box(doc, text):
    # Create a single-cell table for a callout box
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell = table.cell(0, 0)
    
    # Set background color to Isabelline (#F8F2ED)
    set_cell_background(cell, "F8F2ED")
    
    # Set left border to thick Black (#192027) - consistent with Neo-brutalism style
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = parse_xml(r'''
        <w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
            <w:top w:val="none"/>
            <w:left w:val="single" w:sz="24" w:space="0" w:color="192027"/>
            <w:bottom w:val="none"/>
            <w:right w:val="none"/>
        </w:tcBorders>
    ''')
    tcPr.append(tcBorders)
    
    # Add text
    p = cell.paragraphs[0]
    p.paragraph_format.left_indent = Inches(0.1)
    p.paragraph_format.right_indent = Inches(0.1)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    
    run = p.add_run(text)
    run.font.name = 'Inter'
    run.font.size = Pt(10.5)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x19, 0x20, 0x27) # Black (#192027)

def convert_md_to_docx(md_path, docx_path):
    doc = Document()
    
    # Set standard margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Base style setup using Inter font & Black text color
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Inter'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x19, 0x20, 0x27) # Black (#192027)
    
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
        
    in_blockquote = False
    blockquote_text = []
    
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        
        # Handle blockquote
        if line.startswith('>'):
            in_blockquote = True
            text_part = line.lstrip('>').strip()
            blockquote_text.append(text_part)
            i += 1
            continue
        elif in_blockquote and not line.startswith('>') and line != '':
            blockquote_text.append(line)
            i += 1
            continue
        elif in_blockquote and (line == '' or i == len(lines) - 1):
            full_text = " ".join(blockquote_text)
            add_callout_box(doc, full_text)
            blockquote_text = []
            in_blockquote = False
            if line == '':
                i += 1
                continue
        
        # Center-aligned main title HTML style
        if 'align="center"' in line or line.startswith('<h1') or line.startswith('<p align="center">'):
            clean_text = re.sub('<[^<]+?>', '', line)
            if clean_text:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(18)
                p.paragraph_format.space_after = Pt(12)
                run = p.add_run(clean_text)
                run.font.name = 'Inter'
                run.font.size = Pt(20)
                run.font.bold = True
                run.font.color.rgb = RGBColor(0x30, 0x50, 0xA2) # Cyan Cobalt Blue (#3050A2)
            i += 1
            continue

        # Headers
        if line.startswith('# '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(line[2:])
            run.font.name = 'Inter'
            run.font.size = Pt(20)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x30, 0x50, 0xA2) # Cyan Cobalt Blue (#3050A2)
        elif line.startswith('## '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(line[3:])
            run.font.name = 'Inter'
            run.font.size = Pt(16)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x29, 0xB1, 0x98) # Mountain Meadow (#29B198)
        elif line.startswith('### '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(line[4:])
            run.font.name = 'Inter'
            run.font.size = Pt(14)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x30, 0x50, 0xA2) # Cyan Cobalt Blue (#3050A2)
        elif line.startswith('#### '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(line[5:])
            run.font.name = 'Inter'
            run.font.size = Pt(12)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x19, 0x20, 0x27) # Black (#192027)
            
        # Horizontal rule
        elif line == '---':
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(12)
            pBdr = parse_xml(r'<w:pBdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:bottom w:val="single" w:sz="6" w:space="1" w:color="192027"/></w:pBdr>')
            p._p.get_or_add_pPr().append(pBdr)
            
        # Bullet list items
        elif line.startswith('- ') or line.startswith('* '):
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_after = Pt(3)
            text = line[2:]
            parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    r = p.add_run(part[2:-2])
                    r.font.name = 'Inter'
                    r.font.bold = True
                elif part.startswith('*') and part.endswith('*'):
                    r = p.add_run(part[1:-1])
                    r.font.name = 'Inter'
                    r.font.italic = True
                else:
                    r = p.add_run(part)
                    r.font.name = 'Inter'
                    
        # Empty space
        elif line == '':
            pass
            
        # Regular paragraph
        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    r = p.add_run(part[2:-2])
                    r.font.name = 'Inter'
                    r.font.bold = True
                elif part.startswith('*') and part.endswith('*'):
                    r = p.add_run(part[1:-1])
                    r.font.name = 'Inter'
                    r.font.italic = True
                else:
                    r = p.add_run(part)
                    r.font.name = 'Inter'
        i += 1
        
    doc.save(docx_path)
    print("Successfully converted to docx file with branding style.")

if __name__ == '__main__':
    if len(sys.argv) < 3:
        print("Usage: python convert_to_docx.py <input.md> <output.docx>")
        sys.exit(1)
    
    md_path = sys.argv[1]
    docx_path = sys.argv[2]
    convert_md_to_docx(md_path, docx_path)
